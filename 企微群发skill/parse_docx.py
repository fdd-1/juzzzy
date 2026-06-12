#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""解析 docx → 按一级标题拆模块。

输入：
  --input  docx 文件
  --attachments-dir  附件搜索目录（找 PDF/图片，按文件名里出现的模块关键字匹配）

输出：
  exports/{YYYYMMDD}/modules.json
  形如 [{"name": "S5级别", "texts": [...], "attachments": [{"path":..., "type":"file"}]}]

模块切分规则：
  - 标题样式（Heading 1/2）或文本以 "1. S5级别" / "2. S6级别" 这种「数字. 文字」起头都算
  - 模块标题里的「S5/S6/S7/S8」拿来匹配附件文件名

附件类型映射：
  .pdf .docx .pptx .xlsx → file
  .jpg .jpeg .png .gif    → image
  .mp4 .mov               → video
"""
import sys, io, re, json, argparse, datetime as dt
from pathlib import Path

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

from docx import Document  # type: ignore

SCRIPT_DIR = Path(__file__).resolve().parent

IMG_EXT = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"}
VID_EXT = {".mp4", ".mov", ".avi", ".mkv"}
FILE_EXT = {".pdf", ".docx", ".pptx", ".xlsx", ".doc", ".ppt", ".xls", ".zip"}


def log(m): print(m, flush=True)


def classify(p: Path) -> str:
    ext = p.suffix.lower()
    if ext in IMG_EXT:
        return "image"
    if ext in VID_EXT:
        return "video"
    return "file"


HEADING_PREFIX = re.compile(r"^\s*(\d+)[.．、]\s*(.+?)\s*$")


def is_heading(p) -> bool:
    """python-docx Paragraph 是不是标题：看 style.name 含 Heading，或文字符合 '1. xxx' 模式"""
    style = (p.style.name or "").lower() if p.style else ""
    if "heading" in style or "title" in style:
        return True
    # 全文最多 30 字符且匹配「数字. 文字」格式才算
    text = (p.text or "").strip()
    if 1 <= len(text) <= 30 and HEADING_PREFIX.match(text):
        return True
    return False


def clean_heading_text(t: str) -> str:
    m = HEADING_PREFIX.match(t.strip())
    if m:
        return m.group(2).strip()
    return t.strip()


def parse_docx(docx_path: Path):
    """逐段扫描，遇到标题切模块。表格里第二列「物料」忽略（用本地附件目录代替）。"""
    doc = Document(str(docx_path))

    modules = []
    cur = None

    # 用 element 顺序遍历段落 + 表格
    body = doc.element.body
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for child in body.iterchildren():
        if child in para_map:
            p = para_map[child]
            text = (p.text or "").strip()
            if not text:
                continue
            if is_heading(p):
                name = clean_heading_text(text)
                cur = {"name": name, "texts": [], "attachments": []}
                modules.append(cur)
                continue
            if cur is not None:
                cur["texts"].append(text)
            else:
                # 还没遇到第一个标题，跳过前言
                pass
        elif child in table_map:
            t = table_map[child]
            # 表格第一列是话术，第二列是物料；只取第一列文字
            for row in t.rows:
                if not row.cells:
                    continue
                first_cell_text = "\n".join(
                    p.text.strip() for p in row.cells[0].paragraphs if p.text.strip()
                )
                if not first_cell_text:
                    continue
                # 跳过表头行
                if first_cell_text in ("话术", "话术内容", "文案"):
                    continue
                if cur is not None:
                    cur["texts"].append(first_cell_text)
    return modules


def attach_files(modules: list, attachments_dir: Path):
    """按模块名里的关键字（如 S5/S6/S7/S8）匹配 attachments_dir 下文件。"""
    if not attachments_dir.exists():
        log(f"[WARN] 附件目录不存在: {attachments_dir}")
        return
    files = [f for f in attachments_dir.iterdir() if f.is_file()]
    log(f"[INFO] 在 {attachments_dir} 找到 {len(files)} 个附件文件")
    for m in modules:
        # 模块名里抽 'S\d+' 或别的 token
        tokens = re.findall(r"S\d+|[A-Z]\d+|\d+级别", m["name"])
        if not tokens:
            tokens = [m["name"]]
        log(f"  模块「{m['name']}」匹配 token: {tokens}")
        for f in files:
            fname = f.name
            for tk in tokens:
                if tk in fname:
                    m["attachments"].append({
                        "path": str(f.resolve()),
                        "type": classify(f),
                        "filename": fname,
                    })
                    break


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True, help="docx 文件路径")
    ap.add_argument("--attachments-dir", default=None,
                    help="附件搜索目录，默认与 docx 同目录")
    ap.add_argument("--out", default=None,
                    help="输出 modules.json 路径，默认 exports/{YYYYMMDD}/modules.json")
    args = ap.parse_args()

    docx_path = Path(args.input).resolve()
    if not docx_path.exists():
        log(f"[ERROR] 找不到 {docx_path}")
        sys.exit(1)

    attach_dir = Path(args.attachments_dir).resolve() if args.attachments_dir else docx_path.parent

    log(f"[STEP] 解析 docx: {docx_path}")
    modules = parse_docx(docx_path)
    log(f"[STEP] 解析出 {len(modules)} 个模块")
    for m in modules:
        log(f"  - {m['name']}（{len(m['texts'])} 段文字）")

    log(f"[STEP] 在 {attach_dir} 匹配附件...")
    attach_files(modules, attach_dir)

    # 丢掉没有文字的模块（一般是文档大标题）
    before = len(modules)
    modules = [m for m in modules if m["texts"]]
    if before != len(modules):
        log(f"[STEP] 丢弃 {before - len(modules)} 个无文字模块（一般是文档大标题）")

    if args.out:
        out = Path(args.out).resolve()
    else:
        stamp = dt.datetime.now().strftime("%Y%m%d")
        out = SCRIPT_DIR / "exports" / stamp / "modules.json"
    out.parent.mkdir(parents=True, exist_ok=True)

    out.write_text(json.dumps(modules, ensure_ascii=False, indent=2), encoding="utf-8")
    log(f"[OK] 已写入 {out}")
    log("")
    log("[NEXT] 看一眼 modules.json 没问题就跑 create_template.py")


if __name__ == "__main__":
    main()
